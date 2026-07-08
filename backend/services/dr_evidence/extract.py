"""TRACK 24.13 · Document extraction pipeline for the DR Evidence
Intelligence Engine.

Supported inputs
----------------
* ``application/pdf``          – PyMuPDF (fitz) text extraction; detects
  scanned/no-text PDFs and reports ``scanned_pdf_no_text`` honestly.
* ``.xlsx`` / ``.xlsm``        – openpyxl (data-only mode).
* ``.xls``                     – xlrd 2.x (xls-only, still safe).
* ``.csv``                     – Python stdlib ``csv`` with encoding
  fallback and delimiter sniffing.
* ``.docx``                    – python-docx (paragraphs + table cells).
* ``.txt``                     – multi-encoding fallback.
* ``.doc`` legacy binary       – unsupported today; explicit
  ``unsupported`` status with a clean reason so the AI never guesses.

Contract
--------
Every extractor returns an :class:`ExtractionResult` with:

* ``status``           – one of :data:`EXTRACTION_STATUSES`
* ``text``             – lightly normalized text preview (bounded)
* ``rows``             – tabular rows for spreadsheets/CSV (bounded)
* ``page_count``       – for PDFs
* ``sheet_names``      – for spreadsheets
* ``confidence``       – 0.0-1.0
* ``warnings``         – human-readable notes (never used as facts)
* ``ext_meta``         – engine-specific counters (page limits, etc.)

Never raises. Hard failure paths are captured as ``status="failed"``.

Caps (declared here so tests can lock them):

* :data:`MAX_BYTES`          – 25 MB hard cap; larger → ``too_large``
* :data:`MAX_PAGES`          – 60 PDF pages; excess pages → truncated
* :data:`MAX_ROWS`           – 500 spreadsheet/CSV rows; excess → truncated
* :data:`MAX_TEXT_CHARS`     – 40 000 chars in the final ``text`` field
"""
from __future__ import annotations

import csv
import hashlib
import io
import logging
import re
from dataclasses import asdict, dataclass, field
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


# ── Constants (locked by tests) ─────────────────────────────────────

MAX_BYTES = 25 * 1024 * 1024           # 25 MB
MAX_PAGES = 60
MAX_ROWS = 500
MAX_TEXT_CHARS = 40_000
MAX_ROW_CELLS = 40

EXTRACTION_STATUSES = (
    "not_started",
    "extracted",
    "unsupported",
    "failed",
    "too_large",
    "encrypted",
    "corrupt",
    "scanned_pdf_no_text",
)


# ── Result envelope ─────────────────────────────────────────────────

@dataclass
class ExtractionResult:
    """Canonical extraction envelope."""

    status: str = "not_started"
    text: str = ""
    rows: List[List[str]] = field(default_factory=list)
    page_count: int = 0
    sheet_names: List[str] = field(default_factory=list)
    confidence: float = 0.0
    warnings: List[str] = field(default_factory=list)
    ext_meta: Dict[str, Any] = field(default_factory=dict)
    reason: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


# ── Helpers ─────────────────────────────────────────────────────────

def hash_bytes(data: bytes) -> str:
    """Stable SHA-256 hex for the extraction cache. Prefixed with `sha256:`
    so a caller can distinguish algorithms in the future."""
    return "sha256:" + hashlib.sha256(data).hexdigest()


def _truncate_text(s: str) -> Tuple[str, bool]:
    if len(s) <= MAX_TEXT_CHARS:
        return s, False
    return s[:MAX_TEXT_CHARS] + "\n…[truncated]…", True


def _cap_rows(rows: List[List[str]]) -> Tuple[List[List[str]], bool]:
    if len(rows) <= MAX_ROWS:
        capped_cells = [r[:MAX_ROW_CELLS] for r in rows]
        return capped_cells, False
    capped = [r[:MAX_ROW_CELLS] for r in rows[:MAX_ROWS]]
    return capped, True


def _ext_of(filename: str) -> str:
    _, _, ext = (filename or "").rpartition(".")
    return ("." + ext.lower()) if ext else ""


_WS = re.compile(r"[ \t]+")
_NL = re.compile(r"\n{3,}")


def _normalize_text(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = _WS.sub(" ", s)
    s = _NL.sub("\n\n", s)
    return s.strip()


# ── PDF ─────────────────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> ExtractionResult:
    r = ExtractionResult()
    try:
        import fitz  # PyMuPDF  # noqa: PLC0415
    except Exception as e:  # noqa: BLE001
        r.status = "unsupported"
        r.reason = f"pymupdf_import_failed: {e}"
        return r
    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as e:  # noqa: BLE001
        # PyMuPDF raises for corrupt / encrypted PDFs. Distinguish
        # "needs password" so the manifest is honest.
        err = str(e).lower()
        if "password" in err or "encrypted" in err:
            r.status = "encrypted"
            r.reason = "pdf_password_required"
        else:
            r.status = "corrupt"
            r.reason = f"pymupdf_open_failed: {e}"
        return r
    try:
        # Some encrypted PDFs open but require needs_pass check.
        if getattr(doc, "needs_pass", False) or getattr(doc, "is_encrypted", False):
            r.status = "encrypted"
            r.reason = "pdf_needs_password"
            r.page_count = doc.page_count
            doc.close()
            return r

        r.page_count = doc.page_count
        max_page = min(r.page_count, MAX_PAGES)
        chunks: List[str] = []
        for i in range(max_page):
            try:
                page = doc.load_page(i)
                chunks.append(page.get_text("text") or "")
            except Exception as e:  # noqa: BLE001
                r.warnings.append(f"page_{i+1}_extract_failed: {e}")
        doc.close()

        joined = "\n\n".join(c for c in chunks if c).strip()
        if r.page_count > MAX_PAGES:
            r.warnings.append(
                f"pdf_truncated_to_{MAX_PAGES}_pages_of_{r.page_count}",
            )

        if not joined:
            # No embedded text on any of the pages we scanned — this is
            # the classic scanned-PDF signal. We do NOT run OCR here;
            # the manifest surface will surface this to the AI which is
            # instructed not to guess file contents.
            r.status = "scanned_pdf_no_text"
            r.confidence = 0.0
            r.reason = "no_embedded_text"
            return r

        text, truncated = _truncate_text(_normalize_text(joined))
        if truncated:
            r.warnings.append("pdf_text_truncated")
        r.status = "extracted"
        r.text = text
        r.confidence = 0.9
        r.ext_meta = {"engine": "pymupdf", "pages_scanned": max_page}
        return r
    except Exception as e:  # noqa: BLE001
        r.status = "failed"
        r.reason = f"pdf_extract_exception: {e}"
        try:
            doc.close()
        except Exception:  # noqa: BLE001
            pass
        return r


# ── XLSX / XLSM ─────────────────────────────────────────────────────

def _extract_xlsx(data: bytes) -> ExtractionResult:
    r = ExtractionResult()
    try:
        import openpyxl  # noqa: PLC0415
    except Exception as e:  # noqa: BLE001
        r.status = "unsupported"
        r.reason = f"openpyxl_import_failed: {e}"
        return r
    try:
        wb = openpyxl.load_workbook(
            io.BytesIO(data), data_only=True, read_only=True,
        )
    except Exception as e:  # noqa: BLE001
        err = str(e).lower()
        if "password" in err or "encrypted" in err:
            r.status = "encrypted"
            r.reason = "xlsx_encrypted"
        else:
            r.status = "corrupt"
            r.reason = f"openpyxl_open_failed: {e}"
        return r
    try:
        rows: List[List[str]] = []
        r.sheet_names = list(wb.sheetnames)
        total_rows_seen = 0
        row_cap_hit = False
        for sname in r.sheet_names:
            ws = wb[sname]
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                total_rows_seen += 1
                if len(rows) >= MAX_ROWS:
                    row_cap_hit = True
                    break
                cells = [
                    ("" if c is None else str(c)).strip()
                    for c in list(row)[:MAX_ROW_CELLS]
                ]
                # Emit a sheet-name marker at row start so the AI can
                # attribute cells to sheets when useful.
                if i == 0:
                    rows.append([f"[[SHEET:{sname}]]"] + cells)
                else:
                    rows.append(cells)
            if row_cap_hit:
                break
        wb.close()
        if row_cap_hit:
            r.warnings.append(f"xlsx_truncated_to_{MAX_ROWS}_rows")
        rows, _ = _cap_rows(rows)
        r.rows = rows
        # Also emit a light text preview so vector-free consumers see
        # a readable digest.
        text_lines = ["\t".join(row) for row in rows[:100]]
        r.text, _ = _truncate_text(_normalize_text("\n".join(text_lines)))
        r.status = "extracted"
        r.confidence = 0.85
        r.ext_meta = {
            "engine": "openpyxl",
            "sheet_count": len(r.sheet_names),
            "rows_seen": total_rows_seen,
        }
        return r
    except Exception as e:  # noqa: BLE001
        r.status = "failed"
        r.reason = f"xlsx_extract_exception: {e}"
        return r


# ── XLS ─────────────────────────────────────────────────────────────

def _extract_xls(data: bytes) -> ExtractionResult:
    r = ExtractionResult()
    try:
        import xlrd  # noqa: PLC0415
    except Exception as e:  # noqa: BLE001
        r.status = "unsupported"
        r.reason = f"xlrd_import_failed: {e}"
        return r
    try:
        wb = xlrd.open_workbook(file_contents=data)
    except Exception as e:  # noqa: BLE001
        r.status = "corrupt"
        r.reason = f"xls_open_failed: {e}"
        return r
    try:
        rows: List[List[str]] = []
        r.sheet_names = wb.sheet_names()
        row_cap_hit = False
        total = 0
        for sname in r.sheet_names:
            sh = wb.sheet_by_name(sname)
            for i in range(sh.nrows):
                total += 1
                if len(rows) >= MAX_ROWS:
                    row_cap_hit = True
                    break
                vals = [
                    ("" if v is None else str(v)).strip()
                    for v in sh.row_values(i)[:MAX_ROW_CELLS]
                ]
                if i == 0:
                    rows.append([f"[[SHEET:{sname}]]"] + vals)
                else:
                    rows.append(vals)
            if row_cap_hit:
                break
        if row_cap_hit:
            r.warnings.append(f"xls_truncated_to_{MAX_ROWS}_rows")
        rows, _ = _cap_rows(rows)
        r.rows = rows
        r.text, _ = _truncate_text(
            _normalize_text("\n".join("\t".join(row) for row in rows[:100])),
        )
        r.status = "extracted"
        r.confidence = 0.8
        r.ext_meta = {
            "engine": "xlrd", "sheet_count": len(r.sheet_names),
            "rows_seen": total,
        }
        return r
    except Exception as e:  # noqa: BLE001
        r.status = "failed"
        r.reason = f"xls_extract_exception: {e}"
        return r


# ── CSV ─────────────────────────────────────────────────────────────

_CSV_ENCODINGS = ("utf-8-sig", "utf-8", "latin-1", "cp1252")


def _extract_csv(data: bytes) -> ExtractionResult:
    r = ExtractionResult()
    text: Optional[str] = None
    used_enc: Optional[str] = None
    for enc in _CSV_ENCODINGS:
        try:
            text = data.decode(enc)
            used_enc = enc
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        r.status = "corrupt"
        r.reason = "csv_encoding_undecodable"
        return r
    try:
        # Sniff delimiter — safe fallback to comma.
        sample = text[:4096]
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        reader = csv.reader(io.StringIO(text), dialect)
        rows: List[List[str]] = []
        row_cap_hit = False
        total = 0
        for i, row in enumerate(reader):
            total += 1
            if len(rows) >= MAX_ROWS:
                row_cap_hit = True
                break
            rows.append([(c or "").strip() for c in row[:MAX_ROW_CELLS]])
        if row_cap_hit:
            r.warnings.append(f"csv_truncated_to_{MAX_ROWS}_rows")
        rows, _ = _cap_rows(rows)
        r.rows = rows
        r.text, _ = _truncate_text(
            _normalize_text("\n".join("\t".join(row) for row in rows[:200])),
        )
        r.status = "extracted"
        r.confidence = 0.85
        r.ext_meta = {
            "engine": "stdlib_csv",
            "encoding": used_enc,
            "delimiter": getattr(dialect, "delimiter", ","),
            "rows_seen": total,
        }
        return r
    except Exception as e:  # noqa: BLE001
        r.status = "failed"
        r.reason = f"csv_extract_exception: {e}"
        return r


# ── DOCX ────────────────────────────────────────────────────────────

def _extract_docx(data: bytes) -> ExtractionResult:
    r = ExtractionResult()
    try:
        import docx  # noqa: PLC0415  (python-docx)
    except Exception as e:  # noqa: BLE001
        r.status = "unsupported"
        r.reason = f"python_docx_import_failed: {e}"
        return r
    try:
        d = docx.Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        r.status = "corrupt"
        r.reason = f"docx_open_failed: {e}"
        return r
    try:
        parts: List[str] = []
        for p in d.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        # Tables — flatten each row separated by tabs.
        for t in d.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        joined = "\n".join(parts)
        r.text, truncated = _truncate_text(_normalize_text(joined))
        if truncated:
            r.warnings.append("docx_text_truncated")
        r.status = "extracted"
        r.confidence = 0.9
        r.ext_meta = {
            "engine": "python-docx",
            "paragraphs": len(d.paragraphs),
            "tables": len(d.tables),
        }
        return r
    except Exception as e:  # noqa: BLE001
        r.status = "failed"
        r.reason = f"docx_extract_exception: {e}"
        return r


# ── TXT ─────────────────────────────────────────────────────────────

def _extract_txt(data: bytes) -> ExtractionResult:
    r = ExtractionResult()
    for enc in _CSV_ENCODINGS:
        try:
            text = data.decode(enc)
            r.text, truncated = _truncate_text(_normalize_text(text))
            if truncated:
                r.warnings.append("txt_truncated")
            r.status = "extracted"
            r.confidence = 0.95
            r.ext_meta = {"engine": "stdlib_decode", "encoding": enc}
            return r
        except UnicodeDecodeError:
            continue
    r.status = "corrupt"
    r.reason = "txt_encoding_undecodable"
    return r


# ── Dispatcher ──────────────────────────────────────────────────────

def extract_attachment(
    *, filename: str, mime: Optional[str], data: bytes,
) -> ExtractionResult:
    """Route a file to the right extractor.

    Never raises. Every failure mode is captured on the result.
    """
    r = ExtractionResult()
    r.ext_meta["filename"] = filename or ""
    r.ext_meta["mime"] = mime or ""
    r.ext_meta["size_bytes"] = len(data or b"")
    if not data:
        r.status = "failed"
        r.reason = "empty_bytes"
        return r
    if len(data) > MAX_BYTES:
        r.status = "too_large"
        r.reason = f"file_over_{MAX_BYTES}_bytes"
        return r

    ext = _ext_of(filename)
    m = (mime or "").lower()

    if ext == ".pdf" or m == "application/pdf":
        return _extract_pdf(data)
    if ext in (".xlsx", ".xlsm") or m in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel.sheet.macroenabled.12",
    ):
        return _extract_xlsx(data)
    if ext == ".xls" or m == "application/vnd.ms-excel":
        return _extract_xls(data)
    if ext == ".csv" or m in ("text/csv", "application/csv"):
        return _extract_csv(data)
    if ext == ".docx" or m == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return _extract_docx(data)
    if ext == ".txt" or m in ("text/plain",):
        return _extract_txt(data)
    if ext == ".doc" or m == "application/msword":
        r.status = "unsupported"
        r.reason = "legacy_binary_doc_not_supported_use_docx"
        return r
    r.status = "unsupported"
    r.reason = f"no_extractor_for_{ext or m or 'unknown'}"
    return r


__all__ = [
    "EXTRACTION_STATUSES",
    "MAX_BYTES",
    "MAX_PAGES",
    "MAX_ROWS",
    "MAX_ROW_CELLS",
    "MAX_TEXT_CHARS",
    "ExtractionResult",
    "extract_attachment",
    "hash_bytes",
]
