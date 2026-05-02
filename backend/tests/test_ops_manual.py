"""Regression for the ops-manual generator. Confirms both renderers succeed
and produce non-trivial documents with the right top-level structure.
"""
from ops_manual import render_ops_manual_pdf, render_ops_manual_docx


def test_pdf_renders_non_trivial():
    pdf = render_ops_manual_pdf()
    assert len(pdf) > 30_000, f"PDF suspiciously small: {len(pdf)}"
    assert pdf[:4] == b"%PDF"


def test_docx_renders_non_trivial():
    import io
    from docx import Document

    docx_bytes = render_ops_manual_docx()
    assert len(docx_bytes) > 20_000, f"DOCX suspiciously small: {len(docx_bytes)}"
    d = Document(io.BytesIO(docx_bytes))
    # Expect at least one H1 per section (12 sections)
    h1s = [p for p in d.paragraphs if p.style.name == "Heading 1"]
    assert len(h1s) >= 12, f"Expected >=12 H1s, got {len(h1s)}"
    # Expect multiple tables (cost, architecture, failures, etc.)
    assert len(d.tables) >= 8, f"Expected >=8 tables, got {len(d.tables)}"
