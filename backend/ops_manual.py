"""Internal System Owner & Operations Manual — MASCI HUB

Generates both PDF (WeasyPrint) and DOCX (python-docx) for ForgedOps LLC.
Not a customer-facing document. Admin-only download via /api/admin/ops-manual.{pdf,docx}.

Content is kept in structured Python data (SECTIONS) so both renderers emit the
same copy — single source of truth. Tables are rendered as real tables in both
formats (HTML `<table>` for PDF, `doc.add_table` for DOCX).

Re-generated on every request — the document is small (~15-20 pages) so we
don't bother caching.
"""
from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import List, Tuple, Union

from weasyprint import HTML
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Content definition — single source of truth.
#
# Each section is (heading, body_blocks). A body_block is either:
#   • a paragraph string
#   • ("h3", "subheading text")
#   • ("list", [item1, item2, ...])
#   • ("table", [headers...], [[row1_cells], [row2_cells], ...])
# ---------------------------------------------------------------------------

# Type alias for clarity
Block = Union[str, tuple]

SECTIONS: List[Tuple[str, List[Block]]] = [
    ("1. System Overview", [
        "MASCI HUB is the field-operations and safety-documentation platform used by MASCI General Contractors Inc. and MASCI Corporation. It is delivered to end users at the customer-branded URL mascidocs.com. The platform captures, routes, and archives the paperwork that construction crews, mechanics, project managers, and office staff generate every day — Daily Reports, Equipment Pre-Op inspections, Site Safety Inspections, Safety Meetings, JHPs, Incident Reports, Training Packets, and the underlying master data (equipment, employees, suppliers, vendors, jobs, PM assignments).",
        ("h3", "Core Modules"),
        ("table",
            ["Module", "Primary Audience", "What It Does"],
            [
                ["Field Hub", "Foremen, crews, superintendents", "Daily Reports, Equipment Pre-Op, photos, signatures, crew hours, activities"],
                ["Safety Hub", "Safety officers, foremen", "Site inspections, toolbox talks, JHPs, incident reports, trench box tracking, safety cards"],
                ["Shop Hub", "Mechanics, shop supervisors", "Review failed pre-ops, sign off on repairs, parts + work orders"],
                ["PM Hub", "Project managers", "Scoped dashboards per assigned project — daily-report summaries, job-specific incidents, notifications"],
                ["Admin Hub", "Office / ForgedOps", "User roles, master lists (equipment, employees, subs, vendors), integrity checks, backups, training-video URL registry"],
                ["Training Hub", "All personnel (field public, shop/PM/admin gated)", "Step-by-step lessons, bilingual EN/ES content, PDF packets, scan-&-go QR posters"],
            ]),
        ("h3", "Primary Workflows"),
        ("list", [
            "Daily Report — foreman fills at end of shift → auto-routes to assigned PM via email → stored against the job.",
            "Equipment Pre-Op — operator fills at start of shift → FAIL items with photos route to Shop queue → shop signs off before unit returns to service.",
            "Safety Inspection / JHP / Incident — filed as it happens → PM + safety officer notified → archived for OSHA record-keeping.",
            "Training — crews scan a trailer QR code → open /training/field (no login) → watch/read/print the lesson → gated tracks (Shop/PM/Admin) require the matching portal password.",
        ]),
        ("h3", "User Tiers"),
        ("list", [
            "Public (field crews) — no login required for Field/Safety submission forms. Posters use QR codes to deep-link into specific forms.",
            "Shop — per-mechanic email + password account (`shop_users` collection). Access to Shop Hub + Shop training.",
            "PM — per-PM email + password account (`project_managers` collection). Scoped dashboards for assigned projects only.",
            "HR — per-HR-user email + password account (`hr_users` collection). HR portal only — admin tokens do not satisfy /hr/* routes.",
            "Field Leadership — shared password gate (`LEADERSHIP_PASSWORD`). Admin + PM tokens implicitly satisfy.",
            "Admin — email + password account in the `user_directory` collection (bcrypt-12). Full platform control including backups and user management. Legacy single-password gate (`ADMIN_PASSWORD` env) kept as API-only break-glass — no longer reachable from the human-facing UI.",
            "Multi-portal users (iter82) — single email + master password authenticates against /api/auth/multi-login and issues every assigned portal token at once. UI entry point: /sign-in.",
        ]),
    ]),

    ("2. Full System Architecture", [
        ("h3", "Frontend"),
        "React 18 single-page application, bundled with CRACO / Create React App. Tailwind CSS for styling. shadcn/ui component library. React Router for client-side routing. Built output is a ~1.6 MB minified main bundle + code-split chunks, served as static files from the hosting provider. Mobile-first: every form is designed for iPhone SE through desktop. Native HTML5 `<video>` for training videos, not iframes, so field crews on 4G get proper playback controls.",
        ("h3", "Backend"),
        "FastAPI (Python 3.11) running under uvicorn. Single API router at /api prefix. Async endpoints backed by Motor (async MongoDB driver). WeasyPrint for PDF generation. Resend SDK for email. bcrypt for password hashing. JWT for session tokens (signed with JWT_SECRET). Rate limiting and login lockout built in (configurable via env).",
        ("h3", "Database"),
        "MongoDB. Production runs on MongoDB Atlas (managed, free tier sufficient at current scale). Local development and preview environments may use an in-container MongoDB; that instance is ephemeral and wiped on every redeploy — this is why Atlas is mandatory in production.",
        ("h3", "Key Collections"),
        ("table",
            ["Collection", "Purpose"],
            [
                ["user_directory", "Multi-portal master directory — bcrypt-12 hash + portals array (iter82)"],
                ["project_managers", "Per-PM accounts (email + bcrypt) — drives /pm/login"],
                ["shop_users", "Per-mechanic accounts (email + bcrypt) — drives /shop/login"],
                ["hr_users", "Per-HR-user accounts (email + bcrypt) — drives /hr/login"],
                ["admin_audit", "Per-action audit log of multi-portal directory mutations"],
                ["equipment_master", "Fleet registry"],
                ["employees", "Crew roster"],
                ["suppliers / vendors", "Master lists"],
                ["jobs / projects", "Active job registry + PM assignments"],
                ["daily_reports", "Daily report submissions"],
                ["equipment_inspections", "Pre-Op records + FAIL items"],
                ["shop_signoffs", "Shop clearance history"],
                ["incident_reports", "Accidents / near-misses"],
                ["safety_inspections / jhas / meetings", "Safety documentation"],
                ["training_videos", "Config doc: { slug → URL } map"],
                ["training_hits", "Analytics of who/when scanned a poster"],
                ["calculator_runs", "Material Calculator usage analytics (per EN/ES)"],
                ["backup_runs", "Scheduled-backup audit log"],
                ["backup_health", "R2 archive verification ledger (iter79 weekly cron)"],
                ["auth_events", "Login / lockout history"],
            ]),
        ("h3", "File Handling"),
        "User uploads (photos, signatures) are written to Cloudflare R2 object storage (`r2://<bucket>/photos/`, `signatures/`). Database stores only the R2 key — no inline binary blobs after the iter78 migration. Generated PDFs (training packets, records) are rendered in-memory and streamed to the client — not persisted. Local container disk is treated as ephemeral.",
        ("h3", "Email System"),
        "Resend (resend.com) handles all outbound mail — PM routing of daily reports, shop alerts on FAIL pre-ops, bulk safety-card distribution, twice-daily backup emails, and outage alerts. Sender/reply addresses are env-configured (SENDER_EMAIL / REPLY_TO_EMAIL). Emails carry attached PDFs rendered at request time.",
        ("h3", "PDF Generation"),
        "Two generators: (1) /app/backend/pdf_render.py for single records (inspections, reports, incidents); (2) /app/backend/training_pdf.py for multi-page training packets. Both use WeasyPrint with Letter size, custom @page CSS for per-page footer + page count, and a last-page ownership clarification block. Footer text is language-aware: _css_for_lang(lang) injects EN or ES attribution.",
    ]),

    ("3. Third-Party / Supporting Systems", [
        "Every external dependency is listed below. Critical Dependency = impact if the service is down or dropped.",
        ("table",
            ["Service", "Purpose", "Integration", "Criticality"],
            [
                ["Hosting Platform", "Native deployment platform — builds and hosts frontend + backend + MongoDB", "CI-style build on push, serves mascidocs.com + preview URLs", "HIGH — primary runtime"],
                ["MongoDB Atlas", "Production database (managed)", "Backend connects via MONGO_URL env var", "HIGH — all data lives here"],
                ["Cloudflare R2 (S3-compatible)", "Object storage for every photo, signature, and complete-archive backup zip", "boto3 client; env: S3_ENDPOINT_URL, S3_BUCKET, S3_ACCESS_KEY, S3_SECRET_KEY", "HIGH — primary persistence for binary assets and backups"],
                ["Resend", "Transactional email (PM routing, backups, bulk, weekly verification)", "Resend Python SDK, key RESEND_API_KEY", "HIGH — no email = no PM routing, no backups"],
                ["Cloudflare", "CDN + SSL in front of mascidocs.com", "DNS + automatic HTTPS certificate", "HIGH — TLS + caching"],
                ["Domain registrar", "Ownership of mascidocs.com", "Annual renewal, DNS records point at the platform", "HIGH — losing it = losing the brand URL"],
                ["Universal LLM Key", "Access to Gemini/OpenAI/Claude via a single key", "Used for AI-assisted content (translations, future features)", "MEDIUM — optional enrichment, fallback to no-AI path"],
                ["WeasyPrint", "Python PDF renderer (open-source, self-hosted)", "Pinned in requirements.txt", "MEDIUM — generates PDFs; replaceable with wkhtmltopdf if EOL"],
                ["python-docx", "DOCX generator for this manual + future exports", "Pinned in requirements.txt", "LOW — only used for operator docs"],
                ["GitHub (via Save-to-GitHub)", "Source control backup — captures every commit", "Triggered manually from deployment dashboard", "MEDIUM — recovery anchor if the hosting platform is ever abandoned"],
            ]),
    ]),

    ("4. Cost Breakdown", [
        "All figures as of 2026-05. Usage-based costs may rise with crew growth.",
        ("h3", "Per-Service Monthly Cost"),
        ("table",
            ["System", "Purpose", "Monthly Cost (USD)", "Scaling Notes"],
            [
                ["Hosting Standard plan", "Build, host, deploy mascidocs.com + preview", "$20.00", "100 credits/mo included; heavy deploy days can consume 3-5 credits. Pro plan ($200) buys badge-removal + higher limits."],
                ["MongoDB Atlas (M0 Free tier)", "Production database", "$0.00", "M0 = 512 MB storage, shared CPU. Upgrade to M10 (~$57/mo) when storage exceeds 400 MB or crew > 50."],
                ["Resend", "Transactional email", "$0–$20", "Free tier = 3,000 emails/mo. $20/mo tier = 50,000 emails/mo. MASCI should stay under free tier unless bulk safety-card emails exceed 100 recipients weekly."],
                ["Cloudflare (Free)", "CDN, SSL, DNS protection", "$0.00", "Free tier fully covers current traffic. Pro ($20/mo) only needed if image hotlinking or DDoS becomes an issue."],
                ["Domain registrar (mascidocs.com)", "Annual domain renewal", "~$1.20 (annualised)", "$12-15/year typical (.com)."],
                ["Universal LLM Key usage", "AI text via Universal Key", "$0-5", "Pay-per-use. Small model calls are $0.001-0.01 each. Top up via Profile → Universal Key."],
                ["Training video storage", "Customer-asset CDN (platform-hosted)", "$0.00", "Included in hosting platform plan. If video library grows past ~20 videos (~500 MB) consider moving to dedicated S3 + CloudFront."],
                ["Operator time (ForgedOps)", "Admin / maintenance labor", "variable", "1-4 hours/mo at current scale. Rises to 8-10 hours/mo when Daily Reports > 100/week."],
            ]),
        ("h3", "Usage Tier Estimates"),
        ("table",
            ["Tier", "Users", "Jobs", "Reports/Month", "Storage", "Total Monthly"],
            [
                ["Low (current MASCI)", "10-25", "5-10 active", "200-400", "< 300 MB", "$20-40"],
                ["Medium", "25-75", "10-25 active", "500-1,500", "300 MB - 2 GB", "$60-120"],
                ["High", "75-200", "25-80 active", "2,000-6,000", "2-10 GB", "$200-450"],
            ]),
        "Notes: Moving from Low → Medium typically only requires upgrading MongoDB Atlas from M0 free to M10 ($57). Moving Medium → High adds Pro-tier hosting + Resend paid tier + likely S3 migration.",
    ]),

    ("5. Deployment & Environments", [
        ("h3", "Three Environments"),
        ("table",
            ["Environment", "URL", "Purpose"],
            [
                ["Development", "Local container (in-session)", "Live-edit preview while the agent is working — hot-reload, ephemeral DB"],
                ["Preview", "*.preview.emergentagent.com", "Latest committed code, shared database state, visible platform badge"],
                ["Production", "https://mascidocs.com", "Live site for MASCI crews. Atlas-backed, no platform badge"],
            ]),
        ("h3", "How to Deploy"),
        ("list", [
            "Open /admin/system. Check the Pre-Deploy Snapshot panel (iter85) at the top — if GREEN (< 1h old), redeploy is safe. If YELLOW or RED, click 'Snapshot Now' and wait for the build to complete.",
            "Optional: trigger 'BACKUP EVERYTHING' from the Backup & Restore hero panel for an extra redundancy moment (download + email + R2).",
            "Click Save-to-GitHub in the platform chat input. This captures the current frontend + backend as a git commit.",
            "Verify production env vars are set: ADMIN_HMAC_SECRET, JWT_SECRET, SUPER_ADMIN_EMAIL + SUPER_ADMIN_BOOTSTRAP_PASSWORD (for first-boot directory seed), CORS_ORIGINS=https://mascidocs.com, MONGO_URL pointing at Atlas, BACKUP_EMAIL_TO, RESEND_API_KEY, SENDER_EMAIL, REPLY_TO_EMAIL, AUTO_EMAIL_REPORTS=true, RATE_LIMITING=on, BACKUP_R2_HOURLY=true (iter85 — enables hourly cloud snapshots), S3_ENDPOINT_URL + S3_BUCKET + S3_ACCESS_KEY + S3_SECRET_KEY (Cloudflare R2 credentials).",
            "Click Deploy in the deployment dashboard. Build takes 10-15 minutes.",
            "Post-deploy smoke test: curl https://mascidocs.com/api/health → 200. Sign in as Admin via /admin/login or /sign-in (work email + password). Spot-check a dashboard. Confirm BackendVersionBadge in Admin Console footer is GREEN with a new source_hash.",
            "Verify the Pre-Deploy Snapshot panel on /admin/system now shows the deploy as a fresh snapshot baseline.",
        ]),
        ("h3", "Rollback"),
        "Use the Rollback option in the deployment dashboard to return to the previous checkpoint. This reverts code ONLY — if data changed between deploy and rollback, go to /admin/system → Restore from Backup → 'From R2 archive' → pick the snapshot dated just before the bad deploy.",
    ]),

    ("6. Backup & Recovery", [
        ("h3", "Three-Layer Backup Strategy (iter78 + iter79 + iter85)"),
        ("table",
            ["Layer", "Cadence", "Storage", "Purpose"],
            [
                ["Hourly R2 cloud archive (iter85)", "Every UTC hour when BACKUP_R2_HOURLY=true", "Cloudflare R2 — `r2://<bucket>/backups/`", "Closes data-loss window to ~1 hour. Each archive contains Mongo + inlined photo bytes. Self-contained."],
                ["Nightly email backup", "Every UTC night (configurable BACKUP_HOURS_UTC)", "Email attachment to BACKUP_EMAIL_TO", "Off-site copy in admin's inbox. Slim DB + JSON manifest."],
                ["Weekly verification email (iter79)", "Mondays 14:00 UTC", "Email report to BACKUP_VERIFY_TO recipients", "Confirms R2 archives are alive, recent, and well-sized. Positive heartbeat — catches silent upload failures."],
            ]),
        ("h3", "On-Demand Snapshot — Pre-Deploy Snapshot Panel (iter85)"),
        "The /admin/system page now leads with a color-coded freshness widget: GREEN (< 1h old) = safe to redeploy, YELLOW (1-12h) = snapshot recommended, RED (> 12h) = snapshot mandatory. One-click 'Snapshot Now' button fires `POST /api/admin/backups/run-complete-now` which builds a fresh complete archive (~30-60s) and uploads to R2. Always check this panel before any production deploy.",
        ("h3", "Recovery Procedures"),
        ("list", [
            "FASTEST — /admin/system → Restore from Backup → Source = 'From R2 archive' → pick the most recent snapshot → Merge or Replace mode. Restores directly from R2 with no manual download.",
            "EMAIL-BASED — open the most recent nightly backup email → download attachment → /admin/system → Restore from Backup → Source = 'Upload .zip' → pick the file.",
            "WORST-CASE (Atlas + R2 both unavailable) — see Section 9.",
        ]),
        ("h3", "Worst-Case Scenario"),
        "Hosting platform + Atlas + R2 ALL simultaneously unavailable: bring up FastAPI on any Python-hosting provider (Render, Fly.io, Railway), point at a new Atlas cluster (or Docker Mongo), import the most recent nightly backup email via mongorestore, push frontend build to Cloudflare Pages / Vercel / Netlify. Repository is recoverable via GitHub (from Save-to-GitHub commits). Recovery window: 4-8 hours if backups are current. With iter85's hourly R2 cadence, the lost-data window is bounded at ~1 hour even in the catastrophic case.",
    ]),

    ("7. Performance & Scaling", [
        ("h3", "What Impacts Performance"),
        ("table",
            ["Factor", "Typical Cost", "Notes"],
            [
                ["PDF generation (WeasyPrint)", "300-900 ms per record, 2-4 s per training packet", "Synchronous — blocks that worker. Move to a background task queue if concurrent generates exceed 10/min."],
                ["Photo upload", "50-400 ms per image", "Local disk write. Migrating to S3 removes single-host risk but adds ~100 ms network."],
                ["MongoDB query", "1-50 ms", "Acceptable at current volume. Add indexes on (job_id, created_at) when daily_reports > 10k."],
                ["Email send", "300-800 ms", "Resend handles queuing. PM routing is synchronous to the submit request — may be worth offloading to a background task if daily reports > 50/day."],
                ["Cold backend start", "3-6 s", "On platform hibernation policy. First request after idle is slow; subsequent are fast."],
            ]),
        ("h3", "Scaling Behavior"),
        ("list", [
            "More users — no impact up to several hundred concurrent; login throttling (LOGIN_MAX_FAILS / LOGIN_LOCKOUT_SECONDS) is the only bottleneck and is per-IP.",
            "More jobs — MongoDB queries stay under 50 ms as long as indexes are kept. Admin integrity-check may slow down past 10,000 records per collection.",
            "More data (uploads) — local disk fills before anything else breaks. Migrate to S3 before 50 GB total upload volume.",
            "More PDF traffic — WeasyPrint is the chokepoint. Add horizontal workers or a dedicated PDF microservice if > 100 renders/hour.",
        ]),
    ]),

    ("8. Security", [
        ("h3", "Authentication"),
        "Multi-portal master directory (iter82) — `user_directory` collection stores email + bcrypt-12 hash + portals array per user. /api/auth/multi-login authenticates and issues per-portal HMAC tokens in one shot. Per-portal logins (/admin/login, /pm/login, /shop/login, /hr/login) also route through this same endpoint (iter85 admin parity). Legacy single-password gate (ADMIN_PASSWORD env) retained as API-only break-glass — no longer reachable from the human-facing UI. Tokens are stored in browser localStorage and sent via X-Admin-Token / X-PM-Token / X-Shop-Token / X-HR-Token headers depending on portal scope. ADMIN_HMAC_SECRET signs every token; ADMIN_SESSION_EPOCH bump invalidates every issued token in one move.",
        ("h3", "Role-Based Access Control"),
        ("list", [
            "Public — Field/Safety submission forms and Field training only.",
            "Shop — adds Shop Hub, Shop training, sign-off workflows. Per-mechanic accounts in `shop_users`.",
            "PM — adds scoped PM dashboards (only their assigned projects visible). Per-PM accounts in `project_managers`.",
            "HR — adds HR Hub (time verification, accountability, training records). Per-HR-user accounts in `hr_users`. Admin tokens do NOT satisfy /hr/* routes — strict isolation.",
            "Field Leadership — shared password gate. Admin + PM tokens implicitly satisfy.",
            "Admin — full access. Multi-portal directory record with `portals` array including 'admin'. Super-admin flag protects from self-lockout.",
        ]),
        ("h3", "Data Protection"),
        ("list", [
            "TLS everywhere — Cloudflare handles automatic HTTPS for mascidocs.com.",
            "Atlas encrypts data at rest by default.",
            "Passwords hashed, never logged, never returned in API responses.",
            "JWT tokens time-boxed (ACCESS_TOKEN_MINUTES / REFRESH_TOKEN_DAYS configurable).",
            "Rate limiting on public POST endpoints (PUBLIC_POST_LIMIT_PER_HOUR) prevents brute force.",
            "Login lockout after N failures (LOGIN_MAX_FAILS / LOGIN_LOCKOUT_SECONDS).",
            "Admin-only endpoints verify X-Admin-Token HMAC (ADMIN_HMAC_SECRET).",
        ]),
        ("h3", "Client vs Server Exposure"),
        ("list", [
            "Server-only: ADMIN_PASSWORD, PM_PASSWORD, SHOP_PASSWORD, JWT_SECRET, ADMIN_HMAC_SECRET, RESEND_API_KEY, MONGO_URL, EMERGENT_LLM_KEY.",
            "Client-visible (intended): REACT_APP_BACKEND_URL (needed to call the API).",
            "Verified via audit: zero secrets in the minified JS bundle. Grep `sk-`, `re_`, `AWS_`, `MONGO_URL` against the live bundle returns nothing but env-var *names* in admin docs strings.",
        ]),
        ("h3", "Known Risks"),
        ("table",
            ["Risk", "Mitigation"],
            [
                ["Password sharing among crew", "Admin disables the user in /admin/people → directory entry, or rotates their password. Disabled users 401 instantly via password_hash[:16] HMAC binding."],
                ["JWT_SECRET / ADMIN_HMAC_SECRET leak", "Rotate the secret + bump ADMIN_SESSION_EPOCH + redeploy. All tokens instantly invalid across every portal."],
                ["Public endpoint abuse", "Rate limiting already on. Cloudflare adds DDoS absorption."],
                ["Atlas credentials leak", "Atlas IP allowlist + rotate DB user password."],
                ["Cloudflare R2 credentials leak", "Rotate S3_ACCESS_KEY/SECRET in R2 dashboard, update deploy env, redeploy. Old archives stay readable via their original presigned URLs until natural expiry (7 days max)."],
                ["Super admin lockout", "Bootstrap-only re-seed: set SUPER_ADMIN_EMAIL + SUPER_ADMIN_BOOTSTRAP_PASSWORD env vars + restart backend. Directory will recreate the super-admin record if missing."],
                ["Admin audit tampering", "admin_audit collection is append-only via require_admin_strict; mutations to it logged separately. Restore from R2 archive to recover untampered version."],
            ]),
    ]),

    ("9. Failure Points", [
        ("table",
            ["Failure", "What Breaks", "Impact", "Detection", "Fix"],
            [
                ["Hosting platform outage", "Entire platform unreachable", "TOTAL — no one can log in or submit", "Cloudflare 521/522 error, /api/health times out", "Check the hosting platform status page. If extended, trigger DR deploy on alt host per Section 6 worst-case."],
                ["MongoDB Atlas outage", "All reads/writes fail", "TOTAL — app loads but forms error on submit", "Backend logs show connection refused; /api/health 500", "Atlas status page. Atlas usually self-heals. If cluster corrupted, restore from backup ZIP."],
                ["Resend outage", "Emails stop sending (PM routing, backups, alerts)", "HIGH — data still saved; PMs don't get notifications", "No inbound emails for 30+ min; backend logs show Resend 5xx", "Resend status page. All queued data still saved — emails resume when Resend recovers."],
                ["PDF generation failure", "WeasyPrint throws on a specific record", "MEDIUM — one PDF broken; rest of platform fine", "User sees 500 on download; backend logs show WeasyPrint stack trace", "Usually a malformed image upload. Check uploads dir; delete offending file."],
                ["Cloudflare SSL issue", "HTTPS fails, browsers refuse to load", "TOTAL", "Browser 'not secure' warning; curl --insecure still works", "Cloudflare dashboard → SSL/TLS → re-issue cert. Typically auto-heals in <15 min."],
                ["Cloudflare R2 outage", "Photos / signatures / backup downloads fail", "HIGH — record SUBMISSION still works (files queue in app); restore-from-R2 unavailable", "R2 status page; /admin/system Cloud Archives panel shows download failures", "Wait for R2 recovery (typically <30 min). New uploads buffer in memory; older photos served from their R2 cache."],
                ["Domain expiry", "mascidocs.com stops resolving", "TOTAL", "DNS lookup fails", "Renew at registrar. Set auto-renew + 60-day warning."],
                ["Lost super-admin password", "No one can access Admin Console UI", "HIGH — but legacy break-glass API still works", "Sign-in returns 401 with valid email", "Bootstrap re-seed: set SUPER_ADMIN_EMAIL + SUPER_ADMIN_BOOTSTRAP_PASSWORD env vars + restart backend. Directory recreates the super-admin if missing. Alternatively call /api/admin/login (legacy single-pw) via curl with ADMIN_PASSWORD env value as a break-glass."],
                ["Stale Pre-Deploy Snapshot panel (RED state)", "User redeploys without fresh backup — data loss possible", "POTENTIAL — preventive", "Visual: /admin/system top panel is RED", "Click 'Snapshot Now' button. Wait for completion. Then proceed with redeploy."],
                ["Stale backend after deploy (frontend current, backend old)", "New features silently missing from API / PDFs", "MEDIUM", "BackendVersionBadge goes RED or shows old source_hash", "Redeploy."],
                ["User device offline (field crew on bad 4G)", "Can't submit forms", "LOCAL only", "User reports 'it's not saving'", "User retries when signal returns. Consider offline-capable PWA in V2."],
            ]),
    ]),

    ("10. Maintenance Checklist", [
        ("h3", "Daily"),
        ("list", [
            "BackendVersionBadge in Admin Console footer is GREEN (source_hash matches current deploy, uptime sane).",
            "curl https://mascidocs.com/api/health returns 200.",
            "/admin/system → Pre-Deploy Snapshot panel is GREEN (last archive < 1h old).",
            "Spot-check one daily-report email arrived in the PM inbox.",
            "Spot-check one new form submission landed in MongoDB (anything from today).",
        ]),
        ("h3", "Weekly"),
        ("list", [
            "Verify the Monday 14:00 UTC Backup Verification email arrived (iter79 cron) — verdict should be PASS.",
            "Verify nightly backup email is arriving in BACKUP_EMAIL_TO (open one, confirm zip is > 0 bytes).",
            "Glance at /admin/system → Cloud Archives panel — last 7 days of hourly archives should be present.",
            "Review backend error logs for anything repeating (WeasyPrint warnings, Resend 429s, Atlas auth timeouts).",
            "Scan outage_alerts inbox for platform alerts.",
        ]),
        ("h3", "Monthly"),
        ("list", [
            "Review hosting platform credit usage (Profile → Universal Key → Balance). Top up if < 50.",
            "Review Atlas storage usage. Upgrade cluster tier when > 70% of M0 quota.",
            "Review R2 storage usage. With hourly archives, expect ~720 archives/mo (~5-50 GB depending on photo volume). Free tier covers up to 10 GB; consider purging archives > 90 days old if usage climbs.",
            "Review Resend volume. Upgrade tier if > 2,500 emails/mo (80% of free).",
            "Review mascidocs.com domain auto-renew status.",
            "Review new platform features or security advisories (FastAPI, MongoDB, WeasyPrint).",
            "Rotate ADMIN_HMAC_SECRET + JWT_SECRET if > 12 months old; bump ADMIN_SESSION_EPOCH to invalidate every issued token in one move.",
            "Review admin_audit collection for any unexpected directory mutations (admin add/remove, password reset by non-actor).",
        ]),
    ]),

    ("11. Future Scaling / Version 2 Notes", [
        ("h3", "What Should Be Rebuilt for Scale"),
        ("list", [
            "server.py → broken into routers (/routes/auth.py, /routes/training.py, /routes/equipment.py, etc.). Current file is ~9,000 lines — hard to navigate and slows agent-assisted changes. Sub-routers extracted progressively; backup_verification_routes.py is a good template.",
            "Static safety-card PDFs → regenerate via WeasyPrint (same pipeline as training packets) so future footer/logo changes auto-flow to those 4 cards.",
            "PDF generation → background task queue (Celery + Redis, or Arq). Currently blocks the worker; at scale this creates request-timeout risk.",
            "Offline-capable field forms → Progressive Web App with IndexedDB queue. Field crews on bad signal can fill forms offline, sync when online.",
            "Real-time notifications → WebSocket or Server-Sent Events so PMs see daily reports appear without a refresh.",
            "IT Server Dump endpoint — exposes the latest R2 archive's presigned URL via /api/admin/server-dump/latest so MASCI office cron can wget it nightly. P1 on backlog; awaiting IT specs.",
            "Drop the on-disk backup scheduler entirely once IT off-site pulls are running — R2 + verification cron replaces it.",
        ]),
        ("h3", "Dependencies That Could Become Issues"),
        ("list", [
            "Hosting platform lock-in — DR plan mitigates this but worth decoupling if MASCI ever grows to multiple customers.",
            "WeasyPrint — stable but requires GTK system libs; could be replaced with a headless Chrome PDF service (Gotenberg, Browserless) for better CSS support.",
            "Cloudflare R2 free tier (10 GB) — hourly archives accumulate fast at scale. Either lifecycle-purge archives > 30 days OR upgrade to paid R2 (~$0.015/GB/mo).",
        ]),
        ("h3", "Long-Term SaaS Recommendations"),
        ("list", [
            "Separate the MASCI-specific data from the platform code. Introduce a `tenant_id` column on every collection; customer-scope every query.",
            "White-label deployment pattern — each customer gets their own subdomain + isolated Atlas database + configurable branding (logo, lockup, footer text, safety card content).",
            "Per-tenant billing integration (Stripe) so onboarding a new customer doesn't require code changes.",
            "Centralized monitoring (Datadog, Axiom, Better Stack) once more than one customer is on the platform.",
        ]),
    ]),

    ("12. Recent Updates (iter91–iter100, May 2026)", [
        "This section captures architectural and behavioral changes shipped in the latest iteration window. Future iterations should append a new dated subsection rather than overwriting these notes.",
        ("h3", "Admin KPI Strip (iter91–iter93)"),
        "The Admin Console Overview (/admin) leads with a horizontal KPI Strip of 8–10 tiles — each showing lifetime count of a record type plus a 7-day weekly delta arrow plus an optional red alert badge when a queue needs attention. Modules covered: Daily Reports, Pre-Op Inspections, Safety Inspections, Safety Meetings, Incidents, QA/QC, Field Leadership Records, Job Photos, Employees, Equipment. Red badges fire on: Pre-Op FAIL items pending Shop sign-off, unread severe incidents, terminations with outstanding equipment, daily reports flagged by Hours Sanity Flags. Trend math: created_in_last_7_days minus created_in_previous_7_days, recomputed on every page load (no cache).",
        ("h3", "Access Control — Email Delivery Parity (iter90)"),
        "Every action in the Access Control panel (Add User / Reset Password / Disable / Re-enable) now auto-fires a Resend email to the user. Wired identically across PM Accounts, Shop Users, HR Users, and the Multi-Portal Directory. Falls back to backend logs when RESEND_API_KEY is missing or AUTO_EMAIL_REPORTS=false (preview env only). Every action is also written to admin_audit with actor/target/diff/timestamp.",
        ("h3", "Multi-Portal Token Hydration (iter87–iter89)"),
        "Resolved a class of bugs where logging into one portal would silently wipe tokens for other portals on page mount. Three fixes: (1) /api/auth/multi-login provisions shadow records in portal-specific collections (pm_users, shop_users, hr_users) when a user_directory record has portals on it, (2) the frontend token hydrator reconciles localStorage on every route change, (3) the legacy clearToken() calls on login-page mount hooks were removed. Multi-portal users at /sign-in now stay signed in across every entitled portal without re-auth.",
        ("h3", "Uniform Back Buttons (iter96–iter97)"),
        "All `Back to ...` controls across the platform now use a single `<BackLink>` component (/app/frontend/src/components/BackLink.jsx). Role-aware destination (admin → /admin, FL → /leadership, etc.) with consistent visual treatment. Eliminates a long tail of bespoke back-link styles.",
        ("h3", "Termination Email Routing Parity (iter98)"),
        "Employee Termination form (Field Leadership) now auto-CCs the full offboarding loop on submit: assigned PM, HR distribution list, jaymn.judd@mascigc.com, safety@mascigc.com. Subject prefixed `TERMINATION · <Employee> · <Date>`. Law Enforcement flag adds SEVERE_INCIDENT_CC escalation contacts. PDF styling now matches every other Field Leadership form (same letterhead, same `MASCI Operations Platform · Powered by ForgedOps™` footer).",
        ("h3", "FLSA Weekly OT Calculation (iter99)"),
        "HR Time Verification overtime split was reworked to the federal FLSA standard: any hours over 40 in a Mon–Sun week count as overtime. Daily totals are never split into reg/OT — the resolution only happens on the Weekly Rollup view. Regular = min(weekly_total, 40); OT = max(0, weekly_total − 40). Matches Florida construction payroll. Test suite at /app/backend/tests/test_iter99_weekly_ot.py.",
        ("h3", "Hours Sanity Flags (iter100)"),
        "Advisory typo-catcher chips added to NewDailyReport.jsx (per crew row) and HrTimeVerification.jsx (both Weekly Rollup and Per-Day Detail views). Component at /app/frontend/src/components/HoursSanityFlag.jsx. Thresholds: DailyHoursFlag fires on single-day entries > 16h (amber 16.1–24h, red > 24h — catches missing decimals like 60 ≠ 6.0); WeeklyHoursFlag fires on weekly rollups > 80h (amber 80–120h, red > 120h). Flags do NOT block submission — they're advisory chips so foremen catch typos at entry time and HR catches them before payroll cross-check.",
        ("h3", "Time Off Request Workflow (iter102)"),
        "End-to-end employee leave-request system. Two submission paths: (1) Supervisor-submitted via Field Leadership → Time Off Request tile (`/leadership/time_off_request/new`), used when foremen file on behalf of a crew member; (2) Public-link path for office staff — supervisor mints a tokenized URL from the FL hub, employee opens it, submits without any portal login. Both paths land in field_leadership_forms collection with kind=time_off_request and surface in the HR Dashboard for approve/deny. Public-link API: POST /api/field-leadership/time-off/public-link (mints token), POST /api/field-leadership/time-off/public/{link_id} (un-gated public submission). HR routing: assigned PM + HR distribution list + safety@mascigc.com auto-CCed on submission. PDF uses the standardized M-mark letterhead and full footer string. Categories: Vacation · Sick · Medical · Family · Bereavement · Personal.",
        ("h3", "PM Portal Sidebar Architecture (iter105)"),
        "PM Portal restructured to mirror AdminConsole — `PmShell.jsx` is the layout wrapper (amber-600 accent, sticky header w/ M-mark + breadcrumb + portal switcher + health badge + sign-out, collapsible mobile sheet sidebar, 9-section nav). `PmHub.jsx` is now a KPI-tile grid only. Sub-pages live in `pages/pm/PmSections.jsx`: /pm/jobs, /pm/fleet, /pm/people, /pm/suppliers, /pm/posters, /pm/routing, /pm/compliance-export, /pm/field-leadership. The dedicated `/pm/field-leadership` page calls the existing PM-scoped /api/field-leadership endpoint with X-PM-Token — backend already filters records to the PM's assigned jobs server-side, no more re-login prompt that was hitting the password-gated FL SPA.",
        ("h3", "Brand Recalibration (iter104) + Footer Triple-Check (iter105)"),
        "User-facing PDFs and forms now use the M-mark only (no `MASCI HUB` lockup) — `MASCI HUB` is an internal-only name; the platform's external identity is `MASCI Operations Platform`. Every PDF/email/poster footer reads EXACTLY: `GENERATED THROUGH MASCI OPERATIONS PLATFORM — POWERED BY FORGEDOPS™ | © 2026 FORGEDOPS™`. Triple-check pass cleaned residual `MASCI HUB Notifications` From-names and `MASCI HUB · FIELD LEADERSHIP` email body headers in job_photos.py, safety_forms.py, shop_parts.py, and field_leadership.py. Internal-only references (this ops_manual.py, photo_storage.py docstring, outage_alerts.py, server.py admin-backup email subjects) intentionally preserved.",
        ("h3", "Unified Tile UI (iter106–iter108)"),
        "Shared `SectionTile.jsx` component (/app/frontend/src/components/SectionTile.jsx) now powers every tile on the main Hub, the three sub-hubs (FieldSection, SafetySection, QaqcSection), and the Field Leadership Hub. One ACCENTS table (red, redDeep, amber, orange, yellow, lime, emerald, cyan, blue, indigo, purple, fuchsia, rose, slate) drives all tile colors. Visual hierarchy: main hub tiles = headlines only (icon + title + desc + CTA); sub-hub tiles = same anatomy but with section-specific accent. Field Leadership Hub is grouped into 4 logical sections instead of dumped: 01 Daily Crew Documentation · 02 Evaluations & Career Path · 03 Equipment Accountability · 04 HR Actions.",
        ("h3", "Files of reference"),
        ("list", [
            "/app/backend/routes/auth_directory_routes.py — Unified Auth Logic (multi-portal)",
            "/app/frontend/src/components/AdminKpiStrip.jsx — Admin KPI Dashboard",
            "/app/frontend/src/components/BackLink.jsx — Uniform Back Button",
            "/app/frontend/src/components/HoursSanityFlag.jsx — Hours typo catcher",
            "/app/frontend/src/pages/NewDailyReport.jsx — Hours entry (with sanity flag)",
            "/app/frontend/src/pages/HrTimeVerification.jsx — FLSA weekly OT + sanity flags",
            "/app/backend/field_leadership_pdf.py — Standardized FL PDF rendering",
            "/app/backend/routes/hr_portal.py — Weekly OT rollup logic",
        ]),
    ]),

    ("13. Owner Notes", [
        "This section is intentionally blank — ForgedOps™ staff should use this space for ongoing observations, custom changes, and future-improvement ideas as they come up during daily operation of MASCI HUB.",
        ("h3", "Custom Notes"),
        "—",
        ("h3", "Future Improvements"),
        "—",
        ("h3", "Observations"),
        "—",
    ]),
]


COVER_TITLE = "MASCI HUB"
COVER_SUBTITLE = "Internal System Owner & Operations Manual"
COVER_OWNER = "ForgedOps LLC"
COVER_CLASSIFICATION = "CONFIDENTIAL — Not for Customer or Public Use"


# ---------------------------------------------------------------------------
# PDF renderer (WeasyPrint)
# ---------------------------------------------------------------------------

def _pdf_html() -> str:
    now = datetime.now(timezone.utc).strftime("%B %d, %Y")
    parts = []
    parts.append("""<!doctype html><html><head><meta charset="utf-8">
<style>
@page { size: Letter; margin: 0.75in 0.75in 0.95in 0.75in;
  @bottom-left  { content: "MASCI HUB · Internal Operations Manual"; font-family: 'Helvetica', sans-serif; font-size: 8pt; color: #6b7280; }
  @bottom-center { content: "CONFIDENTIAL — ForgedOps LLC"; font-family: 'Helvetica', sans-serif; font-size: 8pt; color: #b91c1c; letter-spacing: 0.1em; text-transform: uppercase; }
  @bottom-right { content: counter(page) " / " counter(pages); font-family: 'Helvetica', sans-serif; font-size: 8pt; color: #6b7280; }
}
@page :first { margin: 1.2in 0.75in 1in 0.75in;
  @bottom-left { content: ""; } @bottom-center { content: ""; } @bottom-right { content: ""; } }
body { font-family: 'Helvetica', Arial, sans-serif; font-size: 10.5pt; color: #1f2937; line-height: 1.5; }
.cover { text-align: center; padding-top: 2in; }
.cover .title { font-size: 32pt; font-weight: 900; letter-spacing: 0.02em; color: #111827; }
.cover .sub { font-size: 14pt; color: #374151; margin-top: 12pt; letter-spacing: 0.05em; }
.cover .owner { margin-top: 40pt; font-size: 11pt; color: #4b5563; font-family: 'Courier New', monospace; text-transform: uppercase; letter-spacing: 0.18em; }
.cover .classification { margin-top: 60pt; color: #b91c1c; font-size: 10pt; font-weight: 700; text-transform: uppercase; letter-spacing: 0.15em; border: 2px solid #b91c1c; padding: 8pt 16pt; display: inline-block; }
.cover .date { margin-top: 28pt; color: #6b7280; font-size: 9pt; font-family: 'Courier New', monospace; }
h1 { font-size: 16pt; font-weight: 900; color: #111827; margin-top: 28pt; margin-bottom: 10pt; padding-bottom: 6pt; border-bottom: 2px solid #111827; page-break-before: always; page-break-after: avoid; }
h3 { font-size: 11pt; font-weight: 700; color: #374151; margin-top: 14pt; margin-bottom: 5pt; text-transform: uppercase; letter-spacing: 0.04em; page-break-after: avoid; }
p { margin: 6pt 0; }
ul { margin: 6pt 0 10pt 18pt; padding: 0; }
li { margin: 3pt 0; }
table { width: 100%; border-collapse: collapse; margin: 8pt 0 12pt; font-size: 9.5pt; }
th { background: #111827; color: white; text-align: left; padding: 6pt 8pt; font-weight: 700; font-size: 9pt; text-transform: uppercase; letter-spacing: 0.04em; }
td { padding: 5pt 8pt; border-bottom: 1px solid #e5e7eb; vertical-align: top; }
tr:nth-child(even) td { background: #f9fafb; }
</style></head><body>""")

    # Cover
    parts.append(f"""<section class="cover">
<div class="title">{COVER_TITLE}</div>
<div class="sub">{COVER_SUBTITLE}</div>
<div class="owner">Prepared for {COVER_OWNER}</div>
<div class="classification">{COVER_CLASSIFICATION}</div>
<div class="date">Generated {now}</div>
</section>""")

    # Sections
    for heading, blocks in SECTIONS:
        parts.append(f"<h1>{_html_escape(heading)}</h1>")
        for b in blocks:
            if isinstance(b, str):
                parts.append(f"<p>{_html_escape(b)}</p>")
            elif b[0] == "h3":
                parts.append(f"<h3>{_html_escape(b[1])}</h3>")
            elif b[0] == "list":
                parts.append("<ul>")
                for item in b[1]:
                    parts.append(f"<li>{_html_escape(item)}</li>")
                parts.append("</ul>")
            elif b[0] == "table":
                headers, rows = b[1], b[2]
                parts.append("<table><thead><tr>")
                for h in headers:
                    parts.append(f"<th>{_html_escape(h)}</th>")
                parts.append("</tr></thead><tbody>")
                for row in rows:
                    parts.append("<tr>")
                    for cell in row:
                        parts.append(f"<td>{_html_escape(cell)}</td>")
                    parts.append("</tr>")
                parts.append("</tbody></table>")

    parts.append("</body></html>")
    return "".join(parts)


def _html_escape(s: str) -> str:
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def render_ops_manual_pdf() -> bytes:
    return HTML(string=_pdf_html()).write_pdf()


# ---------------------------------------------------------------------------
# DOCX renderer (python-docx)
# ---------------------------------------------------------------------------

def render_ops_manual_docx() -> bytes:
    doc = Document()
    # Page margins — 8.5 × 11 is default Letter in python-docx
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n" + COVER_TITLE)
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COVER_SUBTITLE)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\nPrepared for " + COVER_OWNER)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    run.font.name = "Courier New"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n" + COVER_CLASSIFICATION)
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    now = datetime.now(timezone.utc).strftime("%B %d, %Y")
    run = p.add_run("\nGenerated " + now)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.font.name = "Courier New"

    doc.add_page_break()

    # Sections
    for heading, blocks in SECTIONS:
        h = doc.add_heading(heading, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        for b in blocks:
            if isinstance(b, str):
                p = doc.add_paragraph(b)
                p.paragraph_format.space_after = Pt(6)
            elif b[0] == "h3":
                doc.add_heading(b[1], level=3)
            elif b[0] == "list":
                for item in b[1]:
                    doc.add_paragraph(item, style="List Bullet")
            elif b[0] == "table":
                headers, rows = b[1], b[2]
                tbl = doc.add_table(rows=1, cols=len(headers))
                tbl.style = "Light Grid Accent 1"
                hdr_cells = tbl.rows[0].cells
                for i, htext in enumerate(headers):
                    hdr_cells[i].text = htext
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                for row in rows:
                    row_cells = tbl.add_row().cells
                    for i, cell in enumerate(row):
                        row_cells[i].text = str(cell)
                doc.add_paragraph("")  # spacing

    # Footer
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "CONFIDENTIAL — ForgedOps LLC · MASCI HUB Operations Manual"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
